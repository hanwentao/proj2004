import argparse
import csv
import pathlib
import shutil

from django.core.management.base import BaseCommand

import docx
import openpyxl

from ...models import (
    Class,
    Department,
    Profile,
)
from ...templatetags.contacts import gender
from ... import utils


class Command(BaseCommand):
    help = 'Export card information.'

    def add_arguments(self, parser):
        parser.add_argument('-t', '--template',
            type=argparse.FileType('rb'),
            help='Template file in Excel.')
        parser.add_argument('-o', '--output',
            type=argparse.FileType('wb'),
            help='Output file in Excel.')
        parser.add_argument('-p', '--photo',
            help='Path to export photos.')
        parser.add_argument('-H', '--handout', action='store_true',
            help='Export handout materials.')
        parser.add_argument('-O', '--output-path',
            help='Output path for files.')
        parser.add_argument('-m', '--mapping',
            type=argparse.FileType('r'),
            help='Student id to number mapping.')

    def handle(self, *args, **options):
        if options['handout']:
            self.generate_handout(*args, **options)
        else:
            self.generate_handin(*args, **options)

    def generate_handout(self, *args, **options):
        output_path = pathlib.Path(options['output_path'])
        template_path = options['template']
        if options['mapping']:
            mapping = {}
            reader = csv.reader(options['mapping'])
            for sid, num in reader:
                mapping[sid] = num
        else:
            mapping = None
        total = 0
        for department in Department.objects.all():
            department_path = output_path / department.name
            self.stdout.write(f'{department.name}')
            other = False
            for class_ in sorted(department.class_set.all(), key=utils.split_class_name_for_sorted):
                self.stdout.write(f'{class_.name}')
                if utils.split_class_name_for_sorted(class_)[0] == -1:
                    class_path = department_path / f'{class_.name}.docx'
                    doc = docx.Document(template_path)
                    counter = 0
                elif not other:
                    other = True
                    doc = docx.Document(template_path)
                    counter = 0
                table = doc.tables[0]
                for profile in class_.profile_set.filter(user__is_active=True):
                    if profile.class_name != class_.name:  # Not primary class
                        continue
                    user = profile.user
                    extra = user.extra
                    if not extra.photo:
                        continue
                    if not extra.photo_valid:
                        self.stderr.write(str(profile))
                        continue
                    row = table.add_row()
                    counter += 1
                    total += 1
                    if mapping is None:
                        row.cells[0].text = str(counter)
                    else:
                        row.cells[0].text = mapping[profile.student_id]
                    row.cells[1].text = profile.name
                    row.cells[2].text = gender(profile.gender)
                    row.cells[3].text = department.name
                    row.cells[4].text = profile.student_id
                    row.cells[5].text = class_.name
                if utils.split_class_name_for_sorted(class_)[0] == -1 and counter > 0:
                    class_path = department_path / f'{class_.name}.docx'
                    department_path.mkdir(parents=True, exist_ok=True)
                    doc.paragraphs[4].text = doc.paragraphs[4].text.replace('{num_cards}', str(counter))
                    doc.save(class_path)
            if other and counter > 0:
                class_path = department_path / '其他.docx'
                department_path.mkdir(parents=True, exist_ok=True)
                doc.paragraphs[4].text = doc.paragraphs[4].text.replace('{num_cards}', str(counter))
                doc.save(class_path)
        self.stdout.write(str(total))

    def generate_handin(self, *args, **options):
        template_file = options['template']
        output_file = options['output']
        photo_folder_path = pathlib.Path(options['photo']) if options['photo'] else None
        profiles = {}
        num_nonempty = 0
        num_with_photo = 0
        for profile in Profile.objects.filter(user__is_active=True):
            profiles[profile.student_id] = profile
            if profile.completeness > 0:
                num_nonempty += 1
            if profile.user.extra.photo:
                num_with_photo += 1
                if photo_folder_path:
                    photo_path = photo_folder_path / f'{profile.student_id}.jpg'
                    with open(photo_path, 'wb') as photo_file:
                        shutil.copyfileobj(profile.user.extra.photo, photo_file)
        workbook = openpyxl.load_workbook(template_file)
        sheet = workbook.active
        row = 9  # XXX: hacking
        while True:
            student_id = sheet.cell(row=row, column=1).value
            if not student_id:
                break
            student_id = str(student_id)
            profile = profiles.get(student_id)
            if profile is not None:
                sheet.cell(row=row, column=7, value=profile.organization)
                sheet.cell(row=row, column=8, value=' '.join((profile.position, profile.title)))
                sheet.cell(row=row, column=9, value=str(profile.mobile))
                sheet.cell(row=row, column=10, value=profile.email)
                sheet.cell(row=row, column=11, value=profile.address)
                sheet.cell(row=row, column=12, value=profile.postcode)
                sheet.cell(row=row, column=13, value=profile.wechat)
                del profiles[student_id]
            row += 1
        for profile in sorted(profiles.values(), key=lambda p: p.student_id):
            if profile.completeness == 0:
                continue
            sheet.append([
                profile.student_id,
                profile.name,
                gender(profile.gender),
                profile.enroll_year,
                profile.department_name,
                profile.class_name,
                profile.organization,
                ' '.join((profile.position, profile.title)),
                str(profile.mobile),
                profile.email,
                profile.address,
                profile.postcode,
                profile.wechat,
            ])
        workbook.save(output_file)
        self.stdout.write(f'{num_nonempty} {num_with_photo}')
